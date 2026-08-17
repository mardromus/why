"""Extract text from uploaded evidence files for the RAG pipeline."""

from __future__ import annotations

from io import BytesIO
from typing import Any, List, Tuple


SUPPORTED_EXTENSIONS = (".txt", ".md", ".pdf", ".docx", ".json", ".csv")


def _filename(uploaded: Any) -> str:
    return str(getattr(uploaded, "name", "evidence.txt") or "evidence.txt")


def _read_bytes(uploaded: Any) -> bytes:
    if hasattr(uploaded, "getvalue"):
        data = uploaded.getvalue()
        if data:
            return data
    if hasattr(uploaded, "read"):
        data = uploaded.read()
        if hasattr(uploaded, "seek"):
            try:
                uploaded.seek(0)
            except Exception:
                pass
        return data or b""
    if isinstance(uploaded, (bytes, bytearray)):
        return bytes(uploaded)
    return str(uploaded).encode("utf-8", errors="ignore")


def extract_text_from_upload(uploaded: Any) -> Tuple[str, str]:
    """Return (extracted_text, note). Empty text means extraction failed."""
    name = _filename(uploaded)
    lower = name.lower()
    raw = _read_bytes(uploaded)

    if not raw:
        return "", "The uploaded file was empty."

    try:
        if lower.endswith((".txt", ".md", ".csv")):
            return _decode_text(raw), ""
        if lower.endswith(".json"):
            return _decode_text(raw), ""
        if lower.endswith(".pdf"):
            return _extract_pdf(raw)
        if lower.endswith(".docx"):
            return _extract_docx(raw)
        if lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
            return (
                "",
                "Image OCR is not enabled on Streamlit Cloud. Paste the document text or upload PDF/DOCX/TXT.",
            )
        return (
            _decode_text(raw),
            f"Treated '{name}' as plain text. Prefer PDF, DOCX, or TXT for better extraction.",
        )
    except Exception as exc:
        return "", f"Could not extract text from {name}: {exc}"


def extract_text_from_path(path: str) -> Tuple[str, str]:
    with open(path, "rb") as handle:
        class _File:
            name = path
            def getvalue(self_inner):
                return handle.read()

        return extract_text_from_upload(_File())


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore").strip()


def _extract_pdf(raw: bytes) -> Tuple[str, str]:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(raw))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    text = "\n".join(pages).strip()
    if not text:
        return "", "The PDF had no extractable text (it may be a scanned image)."
    return text, ""


def _extract_docx(raw: bytes) -> Tuple[str, str]:
    from docx import Document

    document = Document(BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        return "", "The Word document had no extractable text."
    return text, ""


def chunk_text(text: str, chunk_size: int = 220, overlap: int = 40) -> List[str]:
    words = (text or "").split()
    if not words:
        return []
    if len(words) <= chunk_size:
        return [" ".join(words)]

    chunks = []
    step = max(chunk_size - overlap, 1)
    for start in range(0, len(words), step):
        piece = words[start : start + chunk_size]
        if piece:
            chunks.append(" ".join(piece))
        if start + chunk_size >= len(words):
            break
    return chunks
